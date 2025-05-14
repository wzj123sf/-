from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import tempfile
import os

# 新增的字体设置函数
def set_style_font(style, font_name, east_asian_font_name, font_size=None):
    """设置样式的字体和字号"""
    style.font.name = font_name
    style.element.rPr.rFonts.set(qn('w:eastAsia'), east_asian_font_name)
    if font_size is not None:
        style.font.size = font_size


def create_docx(data):
    doc = Document()

    # 设置全局字体为宋体
    normal_style = doc.styles['Normal']
    set_style_font(normal_style, '宋体', '宋体', Pt(12))

    # 设置标题样式
    heading1 = doc.styles['Heading 1']
    set_style_font(heading1, '宋体', '宋体', Pt(16))

    heading2 = doc.styles['Heading 2']
    set_style_font(heading2, '宋体', '宋体', Pt(14))

    # 标题（应用Heading 1样式）
    title = doc.add_heading(data['metadata']['title'], level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 基本信息（使用Normal样式）
    doc.add_paragraph(f"日期: {data['metadata']['date'][:10]}")
    doc.add_paragraph(f"地点: {data['metadata']['location']}")
    doc.add_paragraph(f"主持人: {data['metadata']['host']['name']}（{data['metadata']['host']['position']}）")

    # 国际形势（应用Heading 2样式）
    doc.add_heading('国际形势分析', level=1)
    doc.add_paragraph(data['agenda']['international']['analysis'])
#C:/Users/邬宗杰/Desktop/meeting_renderer/fonts/simsun.ttc
    # 行动项（应用Heading 2样式）
    doc.add_heading('行动项', level=1)
    table = doc.add_table(rows=1, cols=3)

    # 设置表格字体
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            set_style_font(paragraph.style, '宋体', '宋体')

    # 表格标题行
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '任务'
    hdr_cells[1].text = '负责人'
    hdr_cells[2].text = '截止日期'

    # 添加表格内容
    for item in data['actionItems']:
        if 'task' in item:
            row_cells = table.add_row().cells
            row_cells[0].text = item['task']
            row_cells[1].text = item['owner']
            row_cells[2].text = item['deadline']
            # 设置表格内容字体
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    set_style_font(paragraph.style, '宋体', '宋体')

    # 保存临时文件
    temp_file = tempfile.mktemp(suffix='.docx')
    doc.save(temp_file)
    return temp_file